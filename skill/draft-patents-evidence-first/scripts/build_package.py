#!/usr/bin/env python3
"""Build a release-safe local package with real Markdown, DOCX, PDF, and ZIP files."""
from __future__ import annotations

import argparse
import base64
import hashlib
import html
import json
import mimetypes
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from evidence_first_lib import (
    FULL_LEGAL_DISCLAIMER_EN,
    FULL_LEGAL_DISCLAIMER_ZH,
    ValidationError,
    approved_human_gate,
    command_error,
    load_json,
    now,
    read_case,
    sensitive_findings,
)
from compare_bilingual import validate_bilingual_rows
from inspect_artifacts import InspectorError, inspect, manifest_zip_findings, parse_manifest


def _case_package_members(root: Path) -> list[str]:
    """Return the minimum case record that lets a package validate offline.

    A package is not merely a rendered hand-off.  It must retain the canonical
    configuration and every local source byte named by ``sources.jsonl`` so the
    recipient can re-run validation without reaching back into the source case.
    This helper is deliberately conservative about paths even though ``read_case``
    has already validated them before packaging.
    """
    fixed = [
        "case.yaml",
        "work/sources.jsonl",
        "work/evidence.jsonl",
        "work/claim_trace.jsonl",
        "work/scorecard.json",
        "work/stage_status.json",
    ]
    members = {name for name in fixed if (root / name).is_file()}
    source_index = root / "work" / "sources.jsonl"
    if not source_index.is_file():
        return sorted(members)
    try:
        rows = [json.loads(line) for line in source_index.read_text(encoding="utf-8").splitlines() if line.strip()]
    except (OSError, UnicodeDecodeError, json.JSONDecodeError):
        return sorted(members)
    for row in rows:
        locator = row.get("locator") if isinstance(row, dict) else None
        if not isinstance(locator, str) or locator.startswith("https://"):
            continue
        candidate = Path(locator)
        if candidate.is_absolute() or ".." in candidate.parts:
            continue
        path = root / candidate
        if path.is_file():
            members.add(candidate.as_posix())
    return sorted(members)


def write_manifest(root: Path) -> Path:
    records = []
    for path in sorted(p for p in root.rglob("*") if p.is_file() and p.name != "manifest.json"):
        records.append({
            "path": path.relative_to(root).as_posix(),
            "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
            "bytes": path.stat().st_size,
        })
    manifest = root / "manifest.json"
    manifest.write_text(json.dumps({
        "schema_version": "0.1.0",
        "generated_at": now(),
        "case_package": {
            "portable_case_record": True,
            "canonical_case_files": _case_package_members(root),
        },
        "files": records,
    }, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest


def copy_local_source_material(case_dir: Path, package_dir: Path, sources: list[dict[str, Any]]) -> list[str]:
    """Copy only local bytes referenced by canonical sources into a package.

    ``validate_case.py`` resolves a local locator relative to the case root.  The
    same relative layout is therefore retained in the package.  We reject escaped
    and symlink-resolved paths here as a defence in depth check before copying.
    """
    case_root = case_dir.resolve()
    copied: list[str] = []
    for row in sources:
        locator = row.get("locator")
        if not isinstance(locator, str) or locator.startswith("https://"):
            continue
        relative = Path(locator)
        if relative.is_absolute() or ".." in relative.parts:
            raise ValidationError("local source locator cannot escape case directory")
        source = case_dir / relative
        try:
            source.resolve().relative_to(case_root)
        except ValueError as exc:
            raise ValidationError("local source locator resolves outside case directory") from exc
        if not source.is_file():
            raise ValidationError(f"local source file is unavailable for packaging: {locator}")
        destination = package_dir / relative
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)
        copied.append(relative.as_posix())
    return sorted(set(copied))


def require_real_docx(path: Path) -> None:
    try:
        with zipfile.ZipFile(path) as document:
            if "[Content_Types].xml" not in document.namelist() or "word/document.xml" not in document.namelist():
                raise RuntimeError("generated DOCX is not an OOXML Word document")
    except zipfile.BadZipFile as exc:
        raise RuntimeError("generated DOCX is not a ZIP/OOXML document") from exc


def require_real_pdf(path: Path) -> None:
    if not path.is_file() or not path.read_bytes().startswith(b"%PDF-"):
        raise RuntimeError("generated PDF does not have a PDF signature")


def sanitize_ooxml_metadata(path: Path) -> None:
    """Strip generated application metadata before a package leaves the workspace."""
    with zipfile.ZipFile(path) as source:
        entries = [(item, source.read(item.filename)) for item in source.infolist()]
    with tempfile.NamedTemporaryFile(prefix="efps-docx-", suffix=".docx", dir=path.parent, delete=False) as handle:
        replacement = Path(handle.name)
    try:
        with zipfile.ZipFile(replacement, "w", zipfile.ZIP_DEFLATED) as sanitized:
            for item, payload in entries:
                if item.filename == "docProps/app.xml":
                    payload = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"/>'
                sanitized.writestr(item, payload)
        replacement.replace(path)
    finally:
        replacement.unlink(missing_ok=True)


def markdown_as_html(source: Path, target: Path) -> None:
    """Render the release sample's small Markdown subset for soffice-only hosts."""
    chunks = [
        "<!doctype html><html><head><meta charset=\"utf-8\">",
        "<style>body{font-family:'Noto Sans CJK SC','PingFang SC','Microsoft YaHei',sans-serif;line-height:1.5}",
        "table{border-collapse:collapse;width:100%;table-layout:fixed;margin:12px 0;font-size:9pt}",
        "th,td{border:1px solid #666;padding:5px;text-align:left;vertical-align:top;word-break:break-word;overflow-wrap:anywhere}",
        "img{max-width:90%;height:auto}blockquote{border-left:4px solid #777;padding-left:10px;color:#444}</style></head><body>",
    ]
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    image_pattern = re.compile(r"^!\[(?P<alt>[^]]*)\]\((?P<src>[^)]+)\)$")
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|(?:\s*:?-+:?\s*\|)+\s*$", lines[index + 1]):
            headers = [html.escape(cell.strip()) for cell in stripped.strip("|").split("|")]
            chunks.append("<table><thead><tr>" + "".join(f"<th>{cell}</th>" for cell in headers) + "</tr></thead><tbody>")
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = [html.escape(cell.strip()) for cell in lines[index].strip().strip("|").split("|")]
                chunks.append("<tr>" + "".join(f"<td>{cell}</td>" for cell in cells) + "</tr>")
                index += 1
            chunks.append("</tbody></table>")
            continue
        image = image_pattern.match(stripped)
        if image:
            asset = (source.parent / image.group("src")).resolve()
            if not asset.is_file():
                raise RuntimeError(f"Markdown image is missing: {image.group('src')}")
            mime = mimetypes.guess_type(asset.name)[0] or "application/octet-stream"
            encoded = base64.b64encode(asset.read_bytes()).decode("ascii")
            chunks.append(f'<figure><img src="data:{mime};base64,{encoded}" alt="{html.escape(image.group("alt"))}"></figure>')
        elif stripped.startswith("#"):
            marks = len(stripped) - len(stripped.lstrip("#"))
            text = html.escape(stripped[marks:].strip())
            chunks.append(f"<h{min(marks, 6)}>{text}</h{min(marks, 6)}>")
        elif stripped.startswith(("- ", "* ")):
            chunks.append(f"<p>• {html.escape(stripped[2:])}</p>")
        elif stripped.startswith(">"):
            chunks.append(f"<blockquote>{html.escape(stripped[1:].strip())}</blockquote>")
        else:
            chunks.append(f"<p>{html.escape(stripped)}</p>")
        index += 1
    chunks.append("</body></html>\n")
    target.write_text("\n".join(chunks), encoding="utf-8")


def soffice_convert(source: Path, target: Path, fmt: str, executable: str) -> None:
    with tempfile.TemporaryDirectory(prefix="efps-soffice-") as temp:
        profile = Path(temp) / "profile"
        profile.mkdir()
        result = subprocess.run(
            [
                executable,
                f"-env:UserInstallation={profile.as_uri()}",
                "--headless",
                "--convert-to",
                {"docx": "docx:Office Open XML Text", "pdf": "pdf:writer_pdf_Export"}[fmt],
                "--outdir",
                str(target.parent),
                str(source),
            ],
            text=True,
            capture_output=True,
        )
        produced = target.parent / f"{source.stem}.{fmt}"
        if result.returncode or not produced.exists():
            detail = (result.stderr or result.stdout).strip()
            raise RuntimeError(detail or f"LibreOffice could not convert {source.name} to {fmt}")
        if produced != target:
            produced.replace(target)


def build_docx(source: Path, target: Path, pandoc: str | None, soffice: str | None) -> None:
    if pandoc:
        result = subprocess.run([pandoc, str(source), "-o", str(target)], text=True, capture_output=True)
        if result.returncode:
            raise RuntimeError(result.stderr.strip() or "pandoc DOCX conversion failed")
    elif soffice:
        try:
            from docx import Document
            from docx.enum.section import WD_SECTION
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt
        except ImportError as exc:
            raise RuntimeError("LibreOffice DOCX output also requires python-docx; install requirements-dev.txt") from exc

        document = Document()
        section = document.sections[0]
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(1.8)
        section.left_margin = section.right_margin = Cm(1.8)
        normal = document.styles["Normal"]
        normal.font.name = "Noto Sans CJK SC"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        normal.font.size = Pt(10.5)
        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""

        lines = source.read_text(encoding="utf-8").splitlines()
        index = 0
        image_pattern = re.compile(r"^!\[(?P<alt>[^]]*)\]\((?P<src>[^)]+)\)$")
        while index < len(lines):
            raw = lines[index]
            stripped = raw.strip()
            if not stripped:
                index += 1
                continue
            if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|(?:\s*:?-+:?\s*\|)+\s*$", lines[index + 1]):
                headers = [cell.strip() for cell in stripped.strip("|").split("|")]
                rows: list[list[str]] = []
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                    index += 1
                table = document.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                table.autofit = False
                usable = section.page_width - section.left_margin - section.right_margin
                width = usable // len(headers)
                for cell, value in zip(table.rows[0].cells, headers):
                    cell.width = width
                    cell.text = value
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for values in rows:
                    cells = table.add_row().cells
                    for cell, value in zip(cells, values):
                        cell.width = width
                        cell.text = value
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                run.font.name = "Noto Sans CJK SC"
                                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
                continue
            image = image_pattern.match(stripped)
            if image:
                asset = (source.parent / image.group("src")).resolve()
                if not asset.is_file():
                    raise RuntimeError(f"Markdown image is missing: {image.group('src')}")
                document.add_picture(str(asset), width=Cm(8))
            elif stripped.startswith("#"):
                marks = len(stripped) - len(stripped.lstrip("#"))
                document.add_heading(stripped[marks:].strip(), level=min(marks, 6))
            elif stripped.startswith(("- ", "* ")):
                document.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith(">"):
                document.add_paragraph(stripped[1:].strip(), style="Quote")
            else:
                document.add_paragraph(stripped)
            index += 1
        document.save(target)
    else:
        raise RuntimeError("DOCX/PDF output requires pandoc or LibreOffice/soffice")
    sanitize_ooxml_metadata(target)
    require_real_docx(target)


def build_pdf(source: Path, docx: Path, target: Path, pandoc: str | None, soffice: str | None) -> None:
    if pandoc:
        result = subprocess.run([pandoc, str(source), "-o", str(target)], text=True, capture_output=True)
        if result.returncode and not soffice:
            raise RuntimeError(result.stderr.strip() or "pandoc PDF conversion failed")
        if result.returncode:
            target.unlink(missing_ok=True)
    if not target.exists():
        if not soffice:
            raise RuntimeError("PDF output requires a working pandoc PDF engine or LibreOffice/soffice")
        soffice_convert(docx, target, "pdf", soffice)
    require_real_pdf(target)


def verify_archive(path: Path, package: Path, manifest: Path) -> None:
    try:
        manifest_payload = manifest.read_bytes()
        records = parse_manifest(manifest_payload, str(manifest))
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            internal = [info for info in infos if info.filename == "manifest.json"]
            if len(internal) != 1:
                raise ValidationError("ZIP must contain exactly one manifest.json")
            try:
                archived_manifest = archive.read(internal[0])
            except zipfile.BadZipFile as exc:
                raise ValidationError("ZIP CRC validation failed for manifest.json") from exc
            if archived_manifest != manifest_payload:
                raise ValidationError("ZIP manifest.json differs from package manifest")
            findings = manifest_zip_findings(path, archive, records, internal_manifest=True)
            if findings:
                raise ValidationError("ZIP manifest validation failed: " + json.dumps(findings, ensure_ascii=False, separators=(",", ":")))
            bad = archive.testzip()
            if bad:
                raise ValidationError(f"ZIP CRC validation failed for {bad}")
            for member in infos:
                if not member.filename.isascii() and not member.flag_bits & 0x800:
                    raise ValidationError("ZIP non-ASCII member lacks UTF-8 metadata")
    except InspectorError as exc:
        raise ValidationError(str(exc)) from exc
    except zipfile.BadZipFile as exc:
        raise ValidationError("ZIP is not a readable archive") from exc


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("case_dir", type=Path)
    parser.add_argument("--format", choices=["md", "docx", "pdf", "all"], required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    try:
        # Packaging is an irreversible representation of the local record.  Reuse the
        # complete read-only validator so a scorecard cannot bypass rendered-claim,
        # measured-result, prior-art, or stage-chain reconciliation.
        validator = Path(__file__).with_name("validate_case.py")
        validation = subprocess.run([sys.executable, str(validator), str(args.case_dir)], text=True, capture_output=True)
        if validation.returncode:
            raise ValidationError("case validation blocks packaging: " + (validation.stderr or validation.stdout).strip())
        case, paths, sources, evidence, trace = read_case(args.case_dir)
        # Markdown-only builds are useful while a case is still being drafted.  A final
        # all-format release, however, necessarily includes trace and review artifacts;
        # it cannot represent stages 9–10 as pending/blocked or stale.
        if args.format == "all":
            stages = load_json(paths["stages"])
            incomplete = [str(item.get("stage")) for item in stages if item.get("status") != "complete"]
            if incomplete:
                raise ValidationError("final package requires stages 1–10 complete (incomplete: " + ", ".join(incomplete) + ")")
        if sensitive_findings(args.case_dir):
            raise ValidationError("sensitive findings block packaging")
        gate = case["human_gate"]
        if not approved_human_gate(gate):
            raise ValidationError("approved human selection gate blocks packaging")
        if not trace or any(row["semantic_status"] in {"missing", "conflict"} for row in trace):
            raise ValidationError("missing or conflicting claim limitation trace blocks packaging")
        families = [row.get("family_id") for row in sources if row.get("source_type") == "patent" and row.get("family_id")]
        if len(families) != len(set(families)):
            raise ValidationError("duplicate patent family blocks packaging")
        score = load_json(paths["scorecard"])
        if score.get("status") != "CANDIDATE_CHECKS_PASSED" or score.get("blocking_findings"):
            raise ValidationError("scorecard blocking findings/status block packaging")
        if score.get("unsupported_measured_claims") != 0:
            raise ValidationError("unsupported measured claims block packaging")
        if score.get("duplicate_patent_families") != 0:
            raise ValidationError("scorecard reports duplicate patent families")
        if score.get("claim_trace_coverage_pct") != 100.0 or score.get("evidence_coverage_pct") != 100.0:
            raise ValidationError("scorecard coverage must be 100.0 before packaging")
        if score.get("confidentiality_findings") != 0:
            raise ValidationError("scorecard confidentiality findings block packaging")
        if case["language"] == "bilingual" and score.get("bilingual_atomic_consistency_pct") != 100.0:
            raise ValidationError("bilingual consistency must be 100.0 before packaging")
        if case["language"] == "bilingual":
            bilingual_failures = validate_bilingual_rows(trace)
            if bilingual_failures:
                raise ValidationError("bilingual atomic mismatch blocks packaging: " + json.dumps(bilingual_failures, ensure_ascii=False, separators=(",", ":")))
        if args.output.exists() and any(args.output.iterdir()):
            raise ValidationError("output directory must be empty")

        requested = {"docx", "pdf"} if args.format == "all" else ({args.format} if args.format != "md" else set())
        pandoc, soffice = shutil.which("pandoc"), shutil.which("soffice")
        if requested and not (pandoc or soffice):
            raise RuntimeError("DOCX/PDF output requires pandoc or LibreOffice/soffice")

        args.output.mkdir(parents=True, exist_ok=True)
        work, output = args.output / "work", args.output / "output"
        work.mkdir(); output.mkdir()
        # Keep a portable canonical case root.  ``case.yaml`` and the referenced
        # local source bytes are required for a recipient to validate the extracted
        # ZIP as a real case instead of treating it as an un-runnable report.
        shutil.copy2(paths["case"], args.output / "case.yaml")
        for key in ("sources", "evidence", "trace", "scorecard", "stages"):
            shutil.copy2(paths[key], work / paths[key].name)
        copy_local_source_material(args.case_dir, args.output, sources)
        for source in sorted(paths["output"].rglob("*")):
            if source.is_file():
                destination = output / source.relative_to(paths["output"])
                destination.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(source, destination)
        combined = args.output / "package.md"
        sections = []
        for path in sorted(output.glob("*.md")):
            text = path.read_text(encoding="utf-8")
            text = re.sub(r"(!\[[^]]*\]\()((?![A-Za-z][A-Za-z0-9+.-]*:|/)[^)]+)(\))", r"\1output/\2\3", text)
            sections.append(text)
        combined.write_text("\n\n".join(sections), encoding="utf-8")
        combined_text = combined.read_text(encoding="utf-8")
        if FULL_LEGAL_DISCLAIMER_EN not in combined_text or FULL_LEGAL_DISCLAIMER_ZH not in combined_text:
            raise ValidationError("full bilingual legal disclaimer is required in every generated package")

        docx = args.output / "package.docx"
        if requested:
            build_docx(combined, docx, pandoc, soffice)
            if "docx" not in requested:
                # PDF conversion needs the real DOCX intermediary, but it is not a requested release file.
                docx_for_pdf = docx
            else:
                docx_for_pdf = docx
        else:
            docx_for_pdf = docx
        if "pdf" in requested:
            build_pdf(combined, docx_for_pdf, args.output / "package.pdf", pandoc, soffice)
        if args.format == "pdf":
            docx.unlink(missing_ok=True)

        manifest = write_manifest(args.output)
        archive = args.output.with_suffix(".zip")
        with zipfile.ZipFile(archive, "w", zipfile.ZIP_DEFLATED) as zipped:
            for path in sorted(item for item in args.output.rglob("*") if item.is_file()):
                zipped.write(path, path.relative_to(args.output).as_posix())
        verify_archive(archive, args.output, manifest)
        for artifact in (args.output, archive):
            try:
                findings = inspect(artifact)
            except InspectorError as exc:
                raise RuntimeError(f"artifact inspection unavailable: {exc}") from exc
            if findings:
                raise ValidationError("artifact inspection blocks packaging: " + json.dumps(findings, ensure_ascii=False, separators=(",", ":")))
        print(json.dumps({"package": str(args.output), "manifest": str(manifest), "archive": str(archive)}, ensure_ascii=False))
        return 0
    except RuntimeError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    except ValidationError as exc:
        return command_error(exc)


if __name__ == "__main__":
    sys.exit(main())
